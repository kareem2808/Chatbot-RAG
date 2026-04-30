"""File loader multi-format untuk ingestion ke RAG pipeline.

Mendukung berbagai format file:
- .txt  — Plain text
- .pdf  — PDF documents (via PyPDF2)
- .md   — Markdown files
- .csv  — CSV files (setiap baris = 1 dokumen)
- .json — JSON files (array of strings atau objects dengan key 'text')
- .docx — Word documents (via python-docx, opsional)

Contoh penggunaan:
    >>> from src.core.file_loader import FileLoader
    >>> loader = FileLoader()
    >>> documents = loader.load_file("data/dummy/toko.txt")
    >>> documents = loader.load_directory("data/dummy/")
"""

from __future__ import annotations

import csv
import json
import logging
from dataclasses import dataclass
from io import StringIO
from pathlib import Path
from typing import Any

logger: logging.Logger = logging.getLogger(__name__)

# Extensions yang didukung
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({
    ".txt", ".pdf", ".md", ".csv", ".json", ".docx",
})


@dataclass(frozen=True, slots=True)
class LoadedDocument:
    """Dokumen yang telah di-load dari file.

    Attributes:
        text: Konten teks dokumen.
        metadata: Metadata termasuk sumber file.
    """

    text: str
    metadata: dict[str, Any]


@dataclass
class FileLoader:
    """Loader multi-format untuk membaca file menjadi dokumen teks.

    Otomatis mendeteksi format berdasarkan ekstensi file dan
    mengekstrak teks yang bisa di-ingest ke RAG pipeline.
    """

    def load_file(self, file_path: str | Path) -> list[LoadedDocument]:
        """Load satu file menjadi daftar dokumen.

        Args:
            file_path: Path ke file yang akan di-load.

        Returns:
            Daftar LoadedDocument dari file.

        Raises:
            FileNotFoundError: Jika file tidak ditemukan.
            ValueError: Jika format file tidak didukung.
        """
        path = Path(file_path)
        if not path.exists():
            msg = f"File tidak ditemukan: {path}"
            raise FileNotFoundError(msg)

        ext: str = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            msg = (
                f"Format file '{ext}' tidak didukung. "
                f"Format yang didukung: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )
            raise ValueError(msg)

        base_metadata: dict[str, Any] = {
            "source_file": path.name,
            "source_path": str(path),
            "file_type": ext.lstrip("."),
        }

        logger.info(
            "Memuat file",
            extra={"file": str(path), "type": ext},
        )

        loader_map = {
            ".txt": self._load_txt,
            ".md": self._load_txt,  # Markdown = plain text
            ".pdf": self._load_pdf,
            ".csv": self._load_csv,
            ".json": self._load_json,
            ".docx": self._load_docx,
        }

        loader = loader_map.get(ext)
        if loader is None:
            msg = f"No loader for extension: {ext}"
            raise ValueError(msg)

        documents: list[LoadedDocument] = loader(path, base_metadata)
        logger.info(
            "File berhasil dimuat",
            extra={"file": path.name, "num_documents": len(documents)},
        )
        return documents

    def load_directory(
        self,
        directory: str | Path,
        recursive: bool = True,
    ) -> list[LoadedDocument]:
        """Load semua file yang didukung dari direktori.

        Args:
            directory: Path ke direktori.
            recursive: Jika True, scan subdirektori juga.

        Returns:
            Daftar semua LoadedDocument dari semua file.

        Raises:
            FileNotFoundError: Jika direktori tidak ditemukan.
        """
        dir_path = Path(directory)
        if not dir_path.is_dir():
            msg = f"Direktori tidak ditemukan: {dir_path}"
            raise FileNotFoundError(msg)

        all_documents: list[LoadedDocument] = []
        pattern = "**/*" if recursive else "*"

        for file_path in sorted(dir_path.glob(pattern)):
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                try:
                    docs = self.load_file(file_path)
                    all_documents.extend(docs)
                except Exception as exc:
                    logger.warning(
                        f"Gagal memuat file: {file_path}",
                        extra={"error": str(exc)},
                    )

        logger.info(
            "Direktori berhasil dimuat",
            extra={
                "directory": str(dir_path),
                "total_documents": len(all_documents),
            },
        )
        return all_documents

    # ── Loaders per format ───────────────────
    @staticmethod
    def _load_txt(
        path: Path,
        base_metadata: dict[str, Any],
    ) -> list[LoadedDocument]:
        """Load plain text / markdown file."""
        content: str = path.read_text(encoding="utf-8").strip()
        if not content:
            return []
        return [LoadedDocument(text=content, metadata={**base_metadata})]

    @staticmethod
    def _load_pdf(
        path: Path,
        base_metadata: dict[str, Any],
    ) -> list[LoadedDocument]:
        """Load PDF file — satu dokumen per halaman."""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            msg = (
                "PyPDF2 diperlukan untuk membaca file PDF. "
                "Install: pip install PyPDF2"
            )
            raise ImportError(msg)

        reader = PdfReader(str(path))
        documents: list[LoadedDocument] = []

        for page_num, page in enumerate(reader.pages):
            text: str = page.extract_text() or ""
            text = text.strip()
            if text:
                documents.append(
                    LoadedDocument(
                        text=text,
                        metadata={
                            **base_metadata,
                            "page_number": page_num + 1,
                            "total_pages": len(reader.pages),
                        },
                    )
                )

        return documents

    @staticmethod
    def _load_csv(
        path: Path,
        base_metadata: dict[str, Any],
    ) -> list[LoadedDocument]:
        """Load CSV file — setiap baris menjadi satu dokumen."""
        content: str = path.read_text(encoding="utf-8")
        reader = csv.DictReader(StringIO(content))
        documents: list[LoadedDocument] = []

        for row_num, row in enumerate(reader):
            # Gabungkan semua kolom menjadi teks
            text_parts: list[str] = [
                f"{key}: {value}" for key, value in row.items() if value
            ]
            text: str = "\n".join(text_parts).strip()
            if text:
                documents.append(
                    LoadedDocument(
                        text=text,
                        metadata={
                            **base_metadata,
                            "row_number": row_num + 1,
                        },
                    )
                )

        return documents

    @staticmethod
    def _load_json(
        path: Path,
        base_metadata: dict[str, Any],
    ) -> list[LoadedDocument]:
        """Load JSON file — mendukung array of strings atau objects."""
        content: str = path.read_text(encoding="utf-8")
        data = json.loads(content)
        documents: list[LoadedDocument] = []

        if isinstance(data, list):
            for idx, item in enumerate(data):
                if isinstance(item, str) and item.strip():
                    documents.append(
                        LoadedDocument(
                            text=item.strip(),
                            metadata={**base_metadata, "index": idx},
                        )
                    )
                elif isinstance(item, dict) and "text" in item:
                    text: str = item["text"].strip()
                    extra_meta: dict[str, Any] = {
                        k: v for k, v in item.items() if k != "text"
                    }
                    if text:
                        documents.append(
                            LoadedDocument(
                                text=text,
                                metadata={
                                    **base_metadata,
                                    **extra_meta,
                                    "index": idx,
                                },
                            )
                        )
        elif isinstance(data, dict) and "text" in data:
            text = data["text"].strip()
            if text:
                documents.append(
                    LoadedDocument(text=text, metadata={**base_metadata})
                )

        return documents

    @staticmethod
    def _load_docx(
        path: Path,
        base_metadata: dict[str, Any],
    ) -> list[LoadedDocument]:
        """Load DOCX file (Microsoft Word)."""
        try:
            from docx import Document
        except ImportError:
            msg = (
                "python-docx diperlukan untuk membaca file DOCX. "
                "Install: pip install python-docx"
            )
            raise ImportError(msg)

        doc = Document(str(path))
        paragraphs: list[str] = [
            p.text.strip() for p in doc.paragraphs if p.text.strip()
        ]
        if not paragraphs:
            return []

        full_text: str = "\n\n".join(paragraphs)
        return [LoadedDocument(text=full_text, metadata={**base_metadata})]
